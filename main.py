import os
import logging
import asyncio
import io
from aiogram import Bot, Dispatcher, types
from aiogram.enums import ParseMode, ChatAction
from aiogram.client.default import DefaultBotProperties
from aiogram.filters import Command
from aiogram.types import Message, BufferedInputFile
from docx import Document
import PyPDF2
import requests
from dotenv import load_dotenv

# Загрузка переменных из .env файла
load_dotenv()

# ===== КОНФИГУРАЦИЯ =====
TELEGRAM_BOT_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
YANDEX_API_KEY = os.getenv('YANDEX_API_KEY')
YANDEX_SPEECHKIT_API_KEY = os.getenv('YANDEX_SPEECHKIT_API_KEY')


# Проверка наличия обязательных переменных
if not all([TELEGRAM_BOT_TOKEN, YANDEX_API_KEY, YANDEX_SPEECHKIT_API_KEY]):
    raise ValueError("Не все обязательные переменные окружения заданы в .env файле")

# Инициализация бота
bot = Bot(
    token=TELEGRAM_BOT_TOKEN,
    default=DefaultBotProperties(parse_mode=ParseMode.HTML)
);
dp = Dispatcher()

# Список языков
LANGUAGES = {
    'en': 'Английский',
    'ru': 'Русский',
    'es': 'Испанский',
    'fr': 'Французский',
    'de': 'Немецкий',
    'it': 'Итальянский',
    'ja': 'Японский',
    'zh': 'Китайский',
}

# Максимальные длины
MAX_TEXT_LENGTH = 10000
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

# ===== ФУНКЦИИ ПЕРЕВОДА =====
async def translate_text(text: str, target_lang: str = 'ru') -> str:
    """Перевод текста через Yandex.Translate"""
    if len(text) > MAX_TEXT_LENGTH:
        return f"⚠ Текст слишком длинный (макс. {MAX_TEXT_LENGTH} символов)"

    url = 'https://translate.api.cloud.yandex.net/translate/v2/translate'
    headers = {
        'Authorization': f'Api-Key {YANDEX_API_KEY}',
        'Content-Type': 'application/json',
    }
    body = {
        'targetLanguageCode': target_lang,
        'texts': [text],
    }
    
    try:
        response = requests.post(url, headers=headers, json=body, timeout=10)
        response.raise_for_status()
        return response.json()['translations'][0]['text']
    except Exception as e:
        logging.error(f"Translation error: {e}")
        return f"⚠ Ошибка перевода: {str(e)}"

async def text_to_speech(text: str, lang: str) -> bytes:
    """Генерация аудио через Yandex SpeechKit"""
    if len(text) > 5000:
        return None
        
    url = "https://tts.api.cloud.yandex.net/speech/v1/tts:synthesize"
    headers = {"Authorization": f"Api-Key {YANDEX_SPEECHKIT_API_KEY}"}
    voice = "alena" if lang == "ru" else "john"
    
    data = {
        "text": text,
        "lang": lang,
        "voice": voice,
        "format": "mp3",
        "sampleRateHertz": 48000,
    }
    
    try:
        response = requests.post(url, headers=headers, data=data, timeout=15)
        response.raise_for_status()
        return response.content
    except Exception as e:
        logging.error(f"SpeechKit error: {e}")
        return None

# ===== ОБРАБОТКА ДОКУМЕНТОВ =====
async def process_document(file_info: types.File, file_type: str) -> str:
    """Извлечение текста из документа"""
    try:
        downloaded_file = await bot.download_file(file_info.file_path)
        content = downloaded_file.read()
        
        if file_type == 'docx':
            doc = Document(io.BytesIO(content))
            return '\n'.join([para.text for para in doc.paragraphs])
        elif file_type == 'pdf':
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return '\n'.join([page.extract_text() for page in reader.pages])
        else:  # txt
            return content.decode('utf-8')
    except Exception as e:
        logging.error(f"Document processing error: {e}")
        return None

async def create_translated_file(text: str, original_filename: str) -> BufferedInputFile:
    """Создание файла с переводом"""
    try:
        ext = original_filename.split('.')[-1].lower()
        filename = f"translated_{original_filename}"
        
        if ext == 'docx':
            doc = Document()
            doc.add_paragraph(text)
            buffer = io.BytesIO()
            doc.save(buffer)
            return BufferedInputFile(buffer.getvalue(), filename)
        else:
            return BufferedInputFile(text.encode('utf-8'), filename.replace('.pdf', '.txt'))
    except Exception as e:
        logging.error(f"File creation error: {e}")
        return None

# ===== КОМАНДЫ БОТА =====
@dp.message(Command("start"))
async def start(message: Message):
    await message.answer(
        "Привет! Я бот-переводчик с поддержкой документов.\n\n"
        "Отправь мне текст или файл (txt/docx/pdf) и я переведу его.\n"
        "Примеры:\n"
        "- Текст: \"Hello\" или \"Привет en\"\n"
        "- Файл: прикрепи документ и укажи язык в подписи (например \"en\")\n\n"
        "Команды:\n"
        "/help - справка\n"
        "/langs - список языков"
    )

@dp.message(Command("help"))
async def help(message: Message):
    await message.answer(
        "📌 Как использовать бота:\n\n"
        "1. Для перевода текста:\n"
        "   Просто отправь текст и при необходимости укажи язык через пробел\n"
        "   Пример: \"Hello world es\" - переведет на испанский\n\n"
        "2. Для перевода документов:\n"
        "   Прикрепи файл (txt, docx или pdf) и укажи язык в подписи\n"
        "   Пример: прикрепи файл с подписью \"fr\" - переведет на французский\n\n"
        "3. Максимальный размер файла: 5MB\n"
        "4. Поддерживаемые языки: /langs"
    )

@dp.message(Command("langs"))
async def langs(message: Message):
    langs_list = "\n".join([f"{code} - {name}" for code, name in LANGUAGES.items()])
    await message.answer(f"🌍 Доступные языки:\n{langs_list}")

# ===== ОБРАБОТКА СООБЩЕНИЙ =====
@dp.message(lambda message: message.document)
async def handle_document(message: Message):
    try:
        # Проверка типа и размера файла
        if message.document.file_size > MAX_FILE_SIZE:
            await message.reply("⚠ Файл слишком большой (макс. 5MB)")
            return
            
        file_ext = message.document.file_name.split('.')[-1].lower()
        if file_ext not in ['txt', 'docx', 'pdf']:
            await message.reply("⚠ Поддерживаются только файлы: txt, docx, pdf")
            return

        # Определение языка перевода
        target_lang = 'ru'
        if message.caption and message.caption in LANGUAGES:
            target_lang = message.caption
        
        await bot.send_chat_action(message.chat.id, ChatAction.TYPING)
        progress_msg = await message.reply("⏳ Обрабатываю документ...")

        # Извлечение текста
        file_info = await bot.get_file(message.document.file_id)
        text = await process_document(file_info, file_ext)
        
        if not text:
            await progress_msg.edit_text("⚠ Не удалось извлечь текст из файла")
            return

        # Перевод
        translated_text = await translate_text(text, target_lang)
        if translated_text.startswith("⚠"):
            await progress_msg.edit_text(translated_text)
            return

        # Создание переведенного файла
        translated_file = await create_translated_file(
            translated_text,
            message.document.file_name
        )
        
        if not translated_file:
            await progress_msg.edit_text("⚠ Ошибка создания файла с переводом")
            return

        # Отправка результата
        await message.reply_document(
            document=translated_file,
            caption=f"✅ Перевод на {LANGUAGES.get(target_lang, target_lang)}"
        )
        await progress_msg.delete()

    except Exception as e:
        logging.error(f"Document handling error: {e}")
        await message.reply("⚠ Произошла ошибка при обработке документа")

@dp.message()
async def handle_text(message: Message):
    try:
        user_text = message.text
        target_lang = 'ru'

        # Проверка указания языка
        if ' ' in user_text:
            text_parts = user_text.rsplit(' ', 1)
            if text_parts[1] in LANGUAGES:
                user_text = text_parts[0]
                target_lang = text_parts[1]

        # Проверка длины текста
        if len(user_text) > MAX_TEXT_LENGTH:
            await message.reply(f"⚠ Текст слишком длинный (макс. {MAX_TEXT_LENGTH} символов)")
            return

        await bot.send_chat_action(message.chat.id, ChatAction.TYPING)
        progress_msg = await message.reply("⏳ Обрабатываю запрос...")

        # Перевод
        translated_text = await translate_text(user_text, target_lang)
        if translated_text.startswith("⚠"):
            await progress_msg.edit_text(translated_text)
            return

        # Редактируем сообщение о прогрессе, оставляя перевод
        await progress_msg.edit_text(
            f"🔤 Перевод ({LANGUAGES.get(target_lang, target_lang)}):\n{translated_text}"
        )

        # Озвучка (дополнительно к текстовому переводу)
        audio_data = await text_to_speech(translated_text, target_lang)
        if audio_data:
            await message.reply_voice(
                voice=BufferedInputFile(
                    file=audio_data,
                    filename=f"translation_{target_lang}.mp3"
                ),
                caption="🔊 Озвучка перевода"
            )
        else:
            await message.reply("⚠ Не удалось сгенерировать озвучку (текстовый перевод сохранён выше)")

    except Exception as e:
        logging.error(f"Text handling error: {e}")
        await message.reply("⚠ Произошла ошибка при обработке текста")

# ===== ЗАПУСК БОТА =====
async def main():
    await dp.start_polling(bot)

if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO)
    print("Бот запущен...")
    asyncio.run(main())
